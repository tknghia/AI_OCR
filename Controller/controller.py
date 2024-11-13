import os
import sys
import re
import cv2
import numpy as np
from PIL import Image
from datetime import datetime
from io import BytesIO
import hashlib
import difflib
from pymongo import MongoClient
from bson import ObjectId
from docx import Document
from flask import send_file, jsonify
import matplotlib.pyplot as plt
from pymongo import MongoClient
from bson import ObjectId
import shutil
from sklearn.model_selection import train_test_split
# Add project root to sys.path
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)
sys.path.append(project_root)

# Import vietocr_module and segmentsPaddle from the model
from Model.module import vietocr_module as vietocr_module
from Model.module import crop_text_line_paddle as segmentsPaddle
from Model.module import crop_text_line_yolov8_passport as Passport
from Model.module import crop_text_line_yolov8_driverlicense as DriverLicense
from Model.module import crop_text_line_yolov8_vietnamese_id as VietnameId



class MongoController:
    def __init__(self):
        # MongoDB setup
        MONGO_URI = "mongodb+srv://user:ocrdeeplearning@ocr.83apy.mongodb.net/"
        self.client = MongoClient(MONGO_URI)
        self.db = self.client['OCR_DATABASE']
        self.log_collection = self.db['Logs']
        self.collection_users = self.db['Users']

    def select_prediction_by_id(self, id):
        try:
            object_id = ObjectId(id)  # Convert string ID to ObjectId
            result = self.log_collection.find_one({"_id": object_id})
            return result or f"No prediction found with id: {id}"
        except Exception as e:
            print(f"Error retrieving prediction with id {id}: {e}")
            return f"Error retrieving prediction: {e}"
        
    def get_predictions_by_user_id(self, user_id):
            try:
                # Query to match documents that contain 'user_id' and 'average_accuracy' fields
                query = {
                    "user_id": user_id,
                    "average_accuracy": {"$exists": True}  # Only include logs that have 'average_accuracy' field
                }
                logs = self.log_collection.find(query)
                log_list = list(logs)
                return log_list  # Always return a list, even if empty
            except Exception as e:
                print(f"Error retrieving logs for user with id {user_id}: {e}")
                return []  # Return an empty list in case of error

    def log_predictions(self, predictions, upload_time, list_images, userId=None):
        try:
            mongo_log_entry = {
                "upload_time": upload_time,
                "predictions": predictions,
                "list_images": list_images
            }
            if userId:
                mongo_log_entry["user_id"] = userId

            inserted_id = self.log_collection.insert_one(mongo_log_entry).inserted_id
            return str(inserted_id)

        except Exception as e:
            print(f"Error writing to MongoDB: {e}")
            return None

    def update_labels(self, object_id,comparison_results,user_id=None):
        try:
            query = {"_id": ObjectId(object_id)}
            document = self.log_collection.find_one(query)

            if not document:
                print(f"No document found with id: {object_id}")
                return

            
            avg_levenshtein_similarity = sum(r['levenshtein_similarity'] for r in comparison_results) / len(comparison_results) * 100

            for i, compare in enumerate(comparison_results):
                update_data = {
                    f"list_images.{i}.label": compare["correct"],
                    f"list_images.{i}.prediction": compare["wrong"],
                    f"list_images.{i}.accuracy": compare["levenshtein_similarity"]*100
                }


                # Update label, prediction, and accuracy for each image
                self.log_collection.update_one({"_id": ObjectId(object_id)}, {"$set": update_data})

            # Calculate and update average accuracy
            update_data = {"average_accuracy": avg_levenshtein_similarity}

            # If user_id exists, add it to the update data
            if user_id is not None:
                update_data["user_id"] = user_id

            self.log_collection.update_one({"_id": ObjectId(object_id)}, {"$set": update_data})

            print("Labels, accuracies, average accuracy, and user ID (if provided) updated successfully.")

        except Exception as e:
            print(f"An error occurred: {e}")


    def get_logs_by_user_id(self, user_id):
        try:
            query = {"user_id": user_id}
            logs = self.log_collection.find(query)
            log_list = list(logs)

            return log_list if log_list else f"No logs found for user with id: {user_id}"

        except Exception as e:
            print(f"Error retrieving logs for user with id {user_id}: {e}")
            return f"Error retrieving logs: {e}"
        
    def get_incomplete_accuracy_items(self, document_id):
        try:
            query = {"_id": ObjectId(document_id)}
            document = self.log_collection.find_one(query)

            if not document:
                print(f"No document found with id: {document_id}")
                return []

            upload_time = document.get('upload_time')
            list_images = document.get('list_images', [])
            wrong_labels = document['predictions'].split("\n")

            incomplete_items = [
                {
                    "upload_time": upload_time,
                    "label": image.get('label', ''),
                    "wrong_label": wrong_label,
                    "accuracy": image.get('accuracy', 0)
                }
                for image, wrong_label in zip(list_images, wrong_labels)
                if image.get('accuracy', 100) < 100
            ]

            return [{"upload_time": upload_time, "incomplete_items": incomplete_items}] if incomplete_items else []

        except Exception as e:
            print(f"An error occurred while retrieving items with incomplete accuracy: {e}")
            return []
        

    def update_kyc_info(self, user_id, kyc_info):
        try:
            # Kiểm tra xem user có tồn tại không
            query = {"_id": ObjectId(user_id)}
            user = self.collection_users.find_one(query)

            if not user:
                print(f"No user found with id: {user_id}")
                return {"error": "User not found"}

            # Extract relevant fields from kyc_info
            kyc_dict = {
                "id_card":kyc_info.get("id_card", user.get("id_card", "")),
                "name": kyc_info.get("name", user.get("name", "")),
                "dob": kyc_info.get("dob", user.get("dob", "")),
                "nationality": kyc_info.get("nationality", user.get("nationality", "")),
                "current_place": kyc_info.get("current_place", user.get("current_place", ""))
            }

            # Cập nhật các trường thông tin KYC
            update_fields = {
                "id_card":kyc_dict["id_card"],
                "is_kyc": True,  # Đặt is_kyc là True khi cập nhật thông tin KYC
                "name": kyc_dict["name"],
                "dob": kyc_dict["dob"],
                "nationality": kyc_dict["nationality"],
                "current_place": kyc_dict["current_place"]
            }

            # Cập nhật thông tin trong collection
            self.collection_users.update_one({"_id": ObjectId(user_id)}, {"$set": update_fields})
            print(f"KYC information for user with id {user_id} updated successfully.")
            return {"message": "KYC information updated successfully"}

        except Exception as e:
            print(f"An error occurred while updating KYC information for user with id {user_id}: {e}")
            return {"error": f"An error occurred: {e}"}




class AuthController:
    mongo_controller = MongoController()

    @staticmethod
    def generate_salt():
        # Tạo chuỗi salt ngẫu nhiên (16 bytes)
        return os.urandom(16).hex()

    @staticmethod
    def hash_password(password, salt):
        # Kết hợp password với salt và hash (SHA-256)
        return hashlib.sha256((salt + password).encode()).hexdigest()

    def register(self, data):
        try:
            # Kiểm tra các trường dữ liệu cần thiết
            required_fields = ['username', 'email', 'password']
            if not all(field in data for field in required_fields):
                return {"error": "Missing required fields."}

            username = data['username']
            email = data['email']
            password = data['password']

            # Kiểm tra xem email đã được sử dụng chưa
            if self.mongo_controller.collection_users.find_one({"email": email}):
                return {"error": f"The email {email} is already in use."}

            # Tạo salt và hash password
            salt = self.generate_salt()
            hashed_password = self.hash_password(password, salt)

            # Lưu thông tin user với salt và mật khẩu đã hash
            user_data = {
                "username": username,
                "salt": salt,
                "password": hashed_password,
                "email": email
            }

            self.mongo_controller.collection_users.insert_one(user_data)
            return {"message": "Registration successful."}

        except Exception as e:
            print(f"Error during registration: {e}")
            return {"error": "An error occurred during registration."}

    def login(self, username, password):
        try:
            # Tìm user theo username
            user = self.mongo_controller.collection_users.find_one({"username": username})
            if not user:
                return {"error": "Invalid username or password."}

            # Hash mật khẩu với salt từ user và kiểm tra
            hashed_password = self.hash_password(password, user['salt'])
            if hashed_password == user['password']:
                return {"message": "Login successful.", "user_id": str(user['_id'])}
            else:
                return {"error": "Invalid username or password."}

        except Exception as e:
            print(f"Error during login: {e}")
            return {"error": "An error occurred during login."}
        
    def get_user_info(self, userId):
        user_info = self.mongo_controller.collection_users.find_one(
            {"_id":ObjectId(userId) },
            {"_id": 0, "salt": 0}  
        )

        if user_info:
            return user_info
        else:
            return {"error": "An error occurred during login."}

class ImageController:
    mongo_controller = MongoController()

    @staticmethod
    def reduce_brightness_contrast(image):
        gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced_image = clahe.apply(gray_image)
        return cv2.convertScaleAbs(enhanced_image, alpha=0.8, beta=-30)

    @staticmethod
    def enhance_brightness_contrast(image):
        gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        enhanced_image = clahe.apply(gray_image)
        return cv2.convertScaleAbs(enhanced_image, alpha=1.5, beta=50)

    @staticmethod
    def check_brightness(image):
        gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        return np.mean(gray_image)

    @staticmethod
    def adjust_image_brightness(image):
        brightness_mean = ImageController.check_brightness(image)
        if brightness_mean < 50:
            return ImageController.enhance_brightness_contrast(image)
        elif brightness_mean > 200:
            return ImageController.reduce_brightness_contrast(image)
        return image

    @staticmethod
    def detect_logo_and_rotate(image):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(current_dir, '..', 'View', 'static', 'images', 'logo_cccd.png')
        logo_ref = cv2.imread(logo_path, 0)

        sift = cv2.SIFT_create()
        kp_ref, des_ref = sift.detectAndCompute(logo_ref, None)

        def locate_logo(img):
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img
            kp_img, des_img = sift.detectAndCompute(gray, None)
            if des_img is None:
                return None

            bf = cv2.BFMatcher()
            matches = bf.knnMatch(des_ref, des_img, k=2)
            good_matches = [m for m, n in matches if m.distance < 0.75 * n.distance]

            if len(good_matches) > 5:
                src_pts = np.float32([kp_ref[m.queryIdx].pt for m in good_matches]).reshape(-1, 1, 2)
                dst_pts = np.float32([kp_img[m.trainIdx].pt for m in good_matches]).reshape(-1, 1, 2)
                M, _ = cv2.findHomography(src_pts, dst_pts, cv2.RANSAC, 5.0)
                h, w = logo_ref.shape
                dst = cv2.perspectiveTransform(np.float32([[0, 0], [0, h-1], [w-1, h-1], [w-1, 0]]).reshape(-1, 1, 2), M)
                return np.mean(dst[:, 0, 0]), np.mean(dst[:, 0, 1])
            return None

        def determine_horizontal_orientation(img):
            return img if ImageController.is_wider_than_tall(img) else cv2.rotate(img, cv2.ROTATE_90_CLOCKWISE)

        rotated_image = determine_horizontal_orientation(image)
        logo_position = locate_logo(rotated_image)

        if logo_position:
            height, width = rotated_image.shape[:2]
            center_x, center_y = logo_position
            if center_x < width / 2 and center_y < height / 2:
                return rotated_image
            return cv2.rotate(rotated_image, cv2.ROTATE_180)
        return rotated_image

    @staticmethod
    def is_wider_than_tall(img):
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        if contours:
            x, y, w, h = cv2.boundingRect(max(contours, key=cv2.contourArea))
            return w > h
        return False

    @staticmethod
    def is_face_on_left(img):
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img
        faces = face_cascade.detectMultiScale(gray, scaleFactor=1.05, minNeighbors=15, minSize=(50, 50))
        if not faces:
            return None
        largest_face = max(faces, key=lambda f: f[2] * f[3])
        face_center_x = largest_face[0] + largest_face[2] / 2
        return face_center_x < img.shape[1] / 2

    @staticmethod
    def convert_images(files,type):
            all_predictions = []
            cropped_images_metadata = []

            for file in files:
                img = Image.open(file)
                cv_image = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                rotated_image = ImageController.detect_logo_and_rotate(cv_image)
                enhanced_image = ImageController.adjust_image_brightness(rotated_image)


                height, width = enhanced_image.shape[:2]
                if height > 55:
                    # Perform segmentation and OCR prediction
                    # Kiểm tra loại tài liệu
                    if type == "Passport":
                        # Sử dụng YOLOv8 cho Passport
                        arr = Passport.extract_image_segments(enhanced_image)
                    elif type =="GPLX":
                        arr=DriverLicense.extract_image_segments(enhanced_image)
                    elif type =="CCCD":
                        arr=VietnameId.extract_image_segments(enhanced_image)
                    elif type == "Khác":
                        # Sử dụng PaddleOCR cho các loại khác
                        arr = segmentsPaddle.extract_image_segments(enhanced_image)
                    current_dir = os.path.dirname(os.path.abspath(__file__))
                    output_dir = os.path.join(os.path.dirname(current_dir), 'Model', 'dataset', 'output_images')
                    os.makedirs(output_dir, exist_ok=True)

                    for segment in arr:
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S%f")
                        filename = f"{timestamp}_{segment['id']}.png"
                        filepath = os.path.join(output_dir, filename)
                        Image.fromarray(np.uint8(segment["roi"])).save(filepath)
                        # cropped_images_metadata.append({"filename": filename, "label": None})
                        cropped_images_metadata.append({"filename": filename, "label": segment["type"] if type != "Khác" else None})

                    for segment in arr:
                        image_rgb = cv2.cvtColor(np.asarray(segment["roi"]), cv2.COLOR_BGR2RGB)
                        image_pil = Image.fromarray(image_rgb)
                        prediction = vietocr_module.vietOCR_prediction(image_pil)
                        if type == "Khác":
                            all_predictions.append(f"{prediction}")
                        # Append label and prediction in "label: prediction" format
                        else: 
                            all_predictions.append(f"{segment['type']}: {prediction}")
                else:
                    image_rgb = cv2.cvtColor(cv_image, cv2.COLOR_BGR2RGB)
                    image_pil = Image.fromarray(image_rgb)
                    all_predictions.append(vietocr_module.vietOCR_prediction(image_pil))

            return '\n'.join(all_predictions), cropped_images_metadata

    def process_images(self, files,type,update_profile=None, userId=None):
        upload_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        predictions, cropped_images = self.convert_images(files,type)

        log_file_path = os.path.join(project_root, "Logs", "upload_logs.txt")
        os.makedirs(os.path.dirname(log_file_path), exist_ok=True)

        with open(log_file_path, "a", encoding='utf-8') as log_file:
            log_entries = [f"{upload_time} - Uploaded file: {file.filename}" for file in files]
            log_file.write("\n".join(log_entries) + "\n")
        if update_profile is None:
            prediction_id = self.mongo_controller.log_predictions(predictions, upload_time, cropped_images,userId)
        else:
            prediction_id=None
        return jsonify({'predictions': predictions, 'prediction_id': prediction_id})

    @staticmethod
    def extract_info(data):
        patterns = {
            "Họ và tên": r'Họ và tên / Full name:\s*(.+)',
            "Số căn cước": r'Số / No\.:\s*(\d+)',
            "Ngày sinh": r'Ngày sinh / Date of birth:\s*(\d{2}/\d{2}/\d{4})',
            "Giới tính": r'Giới tính / Sex:\s*(\w+)',
        }
        return {key: re.search(pattern, data).group(1) if re.search(pattern, data) else None for key, pattern in patterns.items()}

class FileController:
    def __init__(self):
        self.mongo_controller = MongoController()

    def save_summary(self, avg_accuracy, sample_count, file_path="summary.docx"):
        try:
            doc = Document(file_path) if os.path.exists(file_path) else Document()
            doc.add_heading(f'Summary of model after {sample_count} samples', level=1)

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'Samples', 'Average Accuracy', 'Loss (Fail)'

            loss_percentage = 100 - avg_accuracy
            row_cells = table.add_row().cells
            row_cells[0].text, row_cells[1].text, row_cells[2].text = str(sample_count), f"{avg_accuracy:.2f}%", f"{loss_percentage:.2f}%"

            doc.save(file_path)
            print(f"Summary saved to {file_path}")
        except Exception as e:
            print(f"An error occurred while saving summary: {e}")

    def calculate_avg_levenshtein_similarity(self, file_path="Logs.docx"):
        try:
            doc = Document(file_path)
            levenshtein_values = [
                float(paragraph.text.split(":")[1].strip().replace("%", ""))
                for paragraph in doc.paragraphs if "Average Levenshtein Similarity:" in paragraph.text
            ]
            avg_similarity = sum(levenshtein_values) / len(levenshtein_values) if levenshtein_values else 0.0
            self.save_summary(avg_similarity, len(levenshtein_values))
            return avg_similarity
        except FileNotFoundError:
            print(f"File {file_path} not found.")
        except Exception as e:
            print(f"An error occurred: {e}")

    def save_comparison_to_word(self, comparison_results, avg_simple_similarity, avg_levenshtein_similarity, file_path):
        doc = Document(file_path) if os.path.exists(file_path) else Document()
        sample_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 1'])
        upload_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        doc.add_heading(f'Samples - {sample_count + 1}', level=1)
        doc.add_heading(f'Upload Time: {upload_time}', level=4)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'Prediction', 'Correct line', 'Simple Similarity', 'Levenshtein Similarity'

        for result in comparison_results:
            row_cells = table.add_row().cells
            row_cells[0].text, row_cells[1].text = str(result['wrong']), str(result['correct'])
            row_cells[2].text, row_cells[3].text = f"{result['simple_similarity']:.2%}", f"{result['levenshtein_similarity']:.2%}"

        doc.add_paragraph(f"Average Simple Similarity: {avg_simple_similarity:.2%}")
        doc.add_paragraph(f"Average Levenshtein Similarity: {avg_levenshtein_similarity:.2%}")
        doc.save(file_path)

        return file_path

    @staticmethod
    def levenshtein_similarity(s1, s2):
        return difflib.SequenceMatcher(None, s1, s2).ratio()

    @staticmethod
    def simple_string_similarity(s1, s2):
        return sum(c1 == c2 for c1, c2 in zip(s1, s2)) / max(len(s1), len(s2))



    @staticmethod
    def normalize_string(s):
        # Loại bỏ ký tự không hiển thị và khoảng trắng thừa
        s = re.sub(r'\s+', ' ', s)  # Thay thế các khoảng trắng liên tiếp bằng 1 khoảng trắng
        return s.strip()

    @classmethod
    def compare_labels(cls, wrong_labels, labels):
        normalized_wrong_labels = [cls.normalize_string(wrong) for wrong in wrong_labels]
        normalized_labels = [cls.normalize_string(correct) for correct in labels]
        return [
            {
                'wrong': wrong,
                'correct': correct,
                'simple_similarity': cls.simple_string_similarity(wrong, correct),
                'levenshtein_similarity': cls.levenshtein_similarity(wrong, correct)
            }
            for wrong, correct in zip(normalized_wrong_labels, normalized_labels)
        ]


    @staticmethod
    def download_content(content):
        buffer = BytesIO()
        buffer.write(content.upper().encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='processed_content.doc', mimetype='application/msword')
    
    def update_file(self, file_path, old_result, labels):
        # Map images to labels and accuracies
        images_labels = {os.path.basename(img["filename"]): label 
                        for img, label in zip(old_result["list_images"], labels)}
        acc_labels = {os.path.basename(img["filename"]): img["accuracy"] 
                    for img in old_result["list_images"]}
        
        updated_lines = []

        if os.path.exists(file_path):
            with open(file_path, "r", encoding='utf-8') as file:
                lines = file.readlines()

            existing_filenames = set()
            for line in lines:
                line = line.strip()
                # Check if line has the correct format, otherwise skip it
                parts = line.split("\t")
                if len(parts) != 2:
                    print(f"Skipping improperly formatted line: {line}")
                    updated_lines.append(line + "\n")
                    continue

                filename, old_label = parts
                if filename in images_labels:
                    # Chỉ cập nhật những dòng có accuracy < 100
                    if acc_labels[filename] < 100:
                        updated_lines.append(f"{filename}\t{images_labels[filename]}\n")
                        existing_filenames.add(filename)
                    else:
                        # Giữ nguyên nhãn cũ nếu accuracy = 100
                        updated_lines.append(line + "\n")
                else:
                    # Keep the line as-is if filename doesn't match
                    updated_lines.append(line + "\n")

            # Add new images that were not in the file and have accuracy < 100
            for filename, label in images_labels.items():
                if filename not in existing_filenames and acc_labels[filename] < 100:
                    updated_lines.append(f"{filename}\t{label}\n")
        else:
            # Create lines from scratch if file doesn't exist, only for accuracy < 100
            updated_lines = [
                f"{os.path.basename(img['filename'])}\t{label}\n"
                for img, label in zip(old_result["list_images"], labels)
                if img["accuracy"] < 100
            ]

        # Write back to the file
        with open(file_path, "w", encoding='utf-8') as file:
            file.writelines(updated_lines)

    def transfer_data(self, datapath):
        """
        Chuyển dữ liệu từ output_dir sang dataset_dir
        nếu file train.txt trong output_dir có >= 100 dòng.
        """
        output_train_path = os.path.join(datapath, 'output_images', 'train.txt')
        output_test_path = os.path.join(datapath, 'output_images', 'test.txt')
        
        # Kiểm tra số dòng trong file train.txt
        with open(output_train_path, 'r', encoding='utf-8') as f:
            num_lines = len(f.readlines())
        
        if num_lines >= 100:

            # Chuyển dữ liệu từ train.txt và test.txt
            dataset_train_path = os.path.join(datapath, 'origin_dataset', 'train.txt')
            dataset_test_path = os.path.join(datapath, 'origin_dataset', 'test.txt')
            
            with open(dataset_train_path, 'a',encoding='utf-8') as dataset_train, \
                open(output_train_path, 'r',encoding='utf-8') as output_train:
                dataset_train.write(output_train.read())
        
            with open(dataset_test_path, 'a',encoding='utf-8') as dataset_test, \
                open(output_test_path, 'r',encoding='utf-8') as output_test:
                dataset_test.write(output_test.read())
            
            # Reset file train.txt và test.txt trong output_dir
            open(output_train_path, 'w', encoding='utf-8').close()
            open(output_test_path, 'w', encoding='utf-8').close()

            # Chuyển ảnh từ output_dir sang dataset_dir
            # Chuyển tất cả file ảnh từ output_images sang origin_dataset
            output_images_dir=os.path.join(datapath, 'output_images')
            origin_dataset_dir=os.path.join(datapath, 'origin_dataset')
            for filename in os.listdir(output_images_dir):
                if filename.endswith('.png') or filename.endswith('.jpg') or filename.endswith('.jpeg'):
                    src = os.path.join(output_images_dir, filename)
                    dst = os.path.join(origin_dataset_dir, filename)
                    shutil.move(src, dst)
    
    def clean_file(self, file_path):
        keywords_to_remove = [
        'current_places:', 'dob:', 'expire_date:', 'gender:', 
        'id:', 'name:', 'nationality:', 'origin_place:',
        'address:', 'code:', 'dob:', 'doi:', 'exp:', 
        'gender:', 'idcard:', 'mrz1:', 'mrz2:', 'name:', 
        'nation:', 'nationality:', 'poi:', 'birth:', 'class:', 
        'day:', 'expiry:', 'month:', 'number:', 'place:', 'year:'
        ]
        
        try:
            # Đọc và xử lý file
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = [line.strip() for line in file if line.strip()]
                cleaned_lines = []
                for line in lines:
                    if '\t' in line:
                        filename, text = line.split('\t', 1)
                        
                        # Xóa các từ khóa khỏi text
                        for keyword in keywords_to_remove:
                            text = text.replace(keyword, '').strip()
                        
                        cleaned_lines.append(f"{filename}\t{text}")
                    else:
                        cleaned_lines.append(line)
            
            # Ghi đè lại file
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write('\n'.join(cleaned_lines))
            
            print(f"Đã xử lý xong file: {file_path}")
            
        except Exception as e:
            print(f"Có lỗi xảy ra: {str(e)}")


    def save_labels(self, labels, result_id, user_id=None):
        log_file_path = os.path.join(project_root, "Model", "dataset", "output_images", "train.txt")
        test_file_path = os.path.join(project_root, "Model", "dataset", "output_images", "test.txt")
        old_result = self.mongo_controller.select_prediction_by_id(result_id)
        # Ensure the folder exists
        os.makedirs(os.path.dirname(log_file_path), exist_ok=True)

        # Process comparison with original predictions
        if not old_result:
            return jsonify({"status": "error", "message": "Invalid result ID"}), 400

        wrong_labels = old_result['predictions'].split("\n")
        comparison_results = self.compare_labels(wrong_labels, labels)
        #check size output images
        datapath=os.path.join(project_root, "Model", "dataset")
        self.transfer_data(datapath)

        # Calculate average correctness percentage
        avg_simple_similarity = sum(r['simple_similarity'] for r in comparison_results) / len(comparison_results)
        avg_levenshtein_similarity = sum(r['levenshtein_similarity'] for r in comparison_results) / len(comparison_results)

        # Save comparison to word document
        self.save_comparison_to_word(comparison_results, avg_simple_similarity, avg_levenshtein_similarity, "Logs.docx")
        self.calculate_avg_levenshtein_similarity("Logs.docx")
        self.mongo_controller.update_labels(result_id,comparison_results, user_id) 
        
        # Get the list of images
        old_result = self.mongo_controller.select_prediction_by_id(result_id)
        images = old_result["list_images"]
        
        train_images, test_images, train_labels, test_labels = train_test_split(images, labels, test_size=0.3, random_state=42)

        # Create temporary results for test and train
        test_result = old_result.copy()
        test_result["list_images"] = test_images
        
        train_result = old_result.copy()
        train_result["list_images"] = train_images
        
        # Update test.txt with first 20% of data
        self.update_file(test_file_path, test_result, test_labels)
        
        # Update train.txt with remaining 80% of data
        self.update_file(log_file_path, train_result, train_labels)
        self.clean_file(log_file_path)
        self.clean_file(test_file_path)
        
        return jsonify({
            "status": "success",
            "avg_simple_similarity": avg_simple_similarity,
            "avg_levenshtein_similarity": avg_levenshtein_similarity
        }), 200
